"""
reporter.py — API Discovery & Security Evaluation Platform
New in this version:
  - Dedicated inbound/outbound API classification section
  - JSON: new top-level key "inbound_outbound_summary" + "outbound_api_inventory"
  - PDF: Section 8 — Inbound vs Outbound API Classification
  - Word: Section 8 — Inbound vs Outbound API Classification
  - outbound_api_inventory read from store.outbound_api_inventory (set by scanner)
  - _dedup_key internal field stripped before output
  - Executive summary updated to mention both inbound and outbound counts
  - console summary shows outbound count
  - All previous fixes preserved
"""

import os
import re
import json
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from store.schema import APIEntry


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _risk_band(score: int) -> str:
    if score >= 75: return "CRITICAL"
    if score >= 50: return "HIGH"
    if score >= 25: return "MEDIUM"
    return "LOW"


def _crit_counts(entries: List[APIEntry]) -> Dict:
    return {
        lvl: sum(1 for e in entries if e.data_sensitivity == lvl)
        for lvl in ["CRITICAL", "HIGH", "MEDIUM", "LOW", "UNKNOWN"]
    }


# ---------------------------------------------------------------------------
# Noise filter
# ---------------------------------------------------------------------------

_ML_EXTENSIONS = re.compile(r"\.(h5|pkl|pt|pth|onnx|pb|model|bin)$", re.I)
_FILE_PATH      = re.compile(r"^(\.\.[\\/]|[A-Za-z]:\\|/home/|/var/|/etc/|/mnt/)")
_MALFORMED_URL  = re.compile(r"^/(https?://)")
_MIDDLEWARE     = re.compile(r"^/?(use|next|req|res|done|:id|:zoneId|:siteId)$", re.I)

def _is_noise(entry: APIEntry) -> bool:
    ep     = entry.endpoint or ""
    method = (entry.method or "").upper()
    if method == "N/A":            return True
    if _FILE_PATH.match(ep):       return True
    if _ML_EXTENSIONS.search(ep):  return True
    if _MALFORMED_URL.match(ep):   return True
    if _MIDDLEWARE.match(ep):      return True
    return False


# ---------------------------------------------------------------------------
# Module / ownership inference
# ---------------------------------------------------------------------------

_MODULE_MAP = [
    (r"vulnerab",                          "Vulnerability Management"),
    (r"login|auth|session|ldap|sso",       "Authentication & Identity"),
    (r"upload|import|ingest",              "Data Upload & Ingestion"),
    (r"menu|permission|role|access|rbac",  "Access Control"),
    (r"itsm|manageengine|ticket|incident", "ITSM Integration"),
    (r"threat|virustotal|malware|ioc",     "Threat Intelligence"),
    (r"crq|quantif|risk",                  "Cyber Risk Quantification"),
    (r"report|export|download",            "Reporting"),
    (r"user|profile|account",             "User Management"),
    (r"device|asset|host",                "Asset Management"),
    (r"dashboard|home|index",             "Dashboard"),
    (r"config|setting|setup",             "Configuration"),
    (r"api/v\d",                           "Public API"),
    (r"media|file|attachment",            "Media & Files"),
    (r"health|ping|status",               "Health & Monitoring"),
    (r"gcp|cloud|storage",                "Cloud Integration"),
]

def _infer_module(entry: APIEntry) -> str:
    file_path = (entry.evidence.get("file", "") if entry.evidence else "").lower()
    endpoint  = (entry.endpoint or "").lower()
    combined  = file_path + " " + endpoint
    for pattern, module in _MODULE_MAP:
        if re.search(pattern, combined):
            return module
    return "Unknown"


def _infer_version(endpoint: str) -> Optional[str]:
    m = re.search(r"/v(\d+)(?:/|$)", endpoint or "")
    return f"v{m.group(1)}" if m else None


# ---------------------------------------------------------------------------
# OWASP flag inference
# ---------------------------------------------------------------------------

def _infer_owasp_flags(entry: APIEntry) -> List[Dict]:
    flags         = list(entry.owasp_flags)
    existing_cats = {f.get("category") for f in flags}
    ep     = entry.endpoint or ""
    method = (entry.method or "").upper()
    auth   = entry.auth_type or "None detected"

    if "API2" not in existing_cats and auth in ("None detected", "UNKNOWN", "none", ""):
        flags.append({"category": "API2", "name": "Broken Authentication",
                      "finding": "No authentication mechanism detected on this endpoint.",
                      "severity": "HIGH", "source": "inferred",
                      "remediation": "Enforce authentication middleware. Use JWT/OAuth2 with short-lived tokens."})

    if "API8" not in existing_cats and ep.startswith("http://"):
        flags.append({"category": "API8", "name": "Security Misconfiguration",
                      "finding": "Endpoint exposed over unencrypted HTTP.",
                      "severity": "MEDIUM", "source": "inferred",
                      "remediation": "Enforce HTTPS. Redirect all HTTP traffic to HTTPS."})

    if "API9" not in existing_cats and entry.classification in ("Shadow", "Rogue"):
        flags.append({"category": "API9", "name": "Improper Inventory Management",
                      "finding": "Endpoint not present in any API registry or specification.",
                      "severity": "MEDIUM", "source": "inferred",
                      "remediation": "Register in API gateway. Create OpenAPI 3.0 spec. Conduct ownership review."})

    if "API9" not in existing_cats and method in ("UNKNOWN", "", "N/A"):
        flags.append({"category": "API9", "name": "Improper Inventory Management",
                      "finding": "HTTP method undocumented — API contract unknown.",
                      "severity": "LOW", "source": "inferred",
                      "remediation": "Document the expected HTTP method and add to OpenAPI spec."})

    if "API1" not in existing_cats and re.search(r"/:\w+|/\{[\w]+\}", ep):
        flags.append({"category": "API1", "name": "Broken Object Level Authorization (BOLA)",
                      "finding": "Parameterized endpoint detected — BOLA risk requires manual testing.",
                      "severity": "HIGH", "source": "inferred",
                      "remediation": "Validate object-level access per user on every request."})

    if "API10" not in existing_cats:
        file_path = (entry.evidence.get("file", "") if entry.evidence else "").lower()
        if any(x in file_path for x in ("virustotal", "manageengine", "external", "third")):
            flags.append({"category": "API10", "name": "Unsafe Consumption of APIs",
                          "finding": "Endpoint integrates with external third-party API.",
                          "severity": "MEDIUM", "source": "inferred",
                          "remediation": "Validate all external API responses. Implement circuit breakers."})

    return flags


# ---------------------------------------------------------------------------
# Outbound dependency extraction (legacy — for backward compat)
# ---------------------------------------------------------------------------

_OUTBOUND_PATTERNS = [
    (r"virustotal",                    "VirusTotal",            "Threat Intelligence",  "External"),
    (r"manageengine|itsm|servicedesk", "ManageEngine ITSM",     "ITSM Integration",     "Internal"),
    (r"ldap",                          "LDAP/Active Directory", "Identity Provider",    "Internal"),
    (r"gcp|googleapis|google.cloud",   "Google Cloud Platform", "Cloud Provider",       "External"),
    (r"threatgraph|threat[\._-]graph", "Threat Graph API",      "Threat Intelligence",  "External"),
    (r"smtp|nodemailer|sendgrid",      "Email Service",         "Notification",         "Internal"),
    (r"kafka|rabbitmq|amqp",           "Message Broker",        "Async Messaging",      "Internal"),
    (r"elasticsearch|opensearch",      "Search Engine",         "Data Store",           "Internal"),
    (r"redis",                         "Redis Cache",           "Cache Layer",          "Internal"),
]

def _extract_outbound_deps(entries: List[APIEntry]) -> List[Dict]:
    """Legacy function — builds simple outbound list from entry text for backward compat."""
    found    = {}
    all_text = []
    for e in entries:
        evidence = e.evidence or {}
        all_text.extend([e.endpoint or "", evidence.get("file", "") or "",
                         evidence.get("match", "") or ""])
        for tag in (e.tags or []):
            all_text.append(str(tag))
    combined = " ".join(all_text).lower()

    for pattern, name, category, exposure in _OUTBOUND_PATTERNS:
        if re.search(pattern, combined) and name not in found:
            found[name] = {
                "integration": name, "category": category, "exposure": exposure,
                "risk": "HIGH" if exposure == "External" else "MEDIUM",
                "recommendation": (
                    f"Validate all data from {name}. Rotate API keys. Implement circuit breakers."
                    if exposure == "External" else
                    f"Ensure {name} uses TLS. Use service account with least-privilege."
                ),
            }
    return list(found.values())


# ---------------------------------------------------------------------------
# Tech stack detection
# ---------------------------------------------------------------------------

def _detect_tech_stack(entries: List[APIEntry], cfg: dict = None) -> Dict:
    scanner_stack = (cfg or {}).get("tech_stack", {})
    cfg_runtime   = scanner_stack.get("runtime", "")
    cfg_framework = scanner_stack.get("framework", "")

    files = []
    for e in entries:
        if e.evidence:
            f = (e.evidence.get("file", "") or "").lower()
            if f:
                files.append(f)
    combined = " ".join(files)

    if cfg_runtime:
        runtime = cfg_runtime
    elif any(f.endswith((".js", ".ts", ".jsx", ".tsx")) for f in files):
        runtime = "Node.js"
    elif any(f.endswith(".py") for f in files):
        runtime = "Python"
    elif any(f.endswith((".java", ".kt")) for f in files):
        runtime = "JVM"
    else:
        runtime = "Unknown"

    if cfg_framework and cfg_framework.lower() not in ("unknown", ""):
        framework = cfg_framework
    else:
        endpoint_sample    = " ".join(e.endpoint or "" for e in entries[:200]).lower()
        has_express_params = bool(re.search(r"/:\w+", endpoint_sample))
        has_express_files  = any("route" in f or "controller" in f or "middleware" in f for f in files)
        if runtime == "Node.js" and (has_express_params or has_express_files):
            framework = "Express.js"
        elif "flask" in combined or (runtime == "Python" and any("route" in f for f in files)):
            framework = "Flask (Python)"
        elif "spring" in combined or runtime == "JVM":
            framework = "Spring Boot (Java)"
        else:
            framework = f"{runtime} (framework not detected)" if runtime != "Unknown" else "Unknown"

    lang_map  = {"Node.js": "JavaScript/TypeScript", "Python": "Python", "JVM": "Java/Kotlin"}
    lang      = lang_map.get(runtime, "Unknown")
    has_react = any("component" in f or f.endswith((".tsx", ".jsx")) for f in files)
    frontend  = "React SPA" if has_react else "Unknown"
    indicators = []
    if framework == "Express.js":
        indicators.append("Express-style path parameters detected (/:id patterns)")
    if has_react:
        indicators.append("React component files (.tsx/.jsx) detected in source")
    if cfg_framework:
        indicators.append(f"Framework confirmed by scanner: {cfg_framework}")

    return {"runtime": runtime, "language": lang, "framework": framework,
            "frontend": frontend, "indicators": indicators}


# ---------------------------------------------------------------------------
# OWASP conformance
# ---------------------------------------------------------------------------

def _build_owasp_conformance(entries: List[APIEntry], all_flags: List[Dict]) -> List[Dict]:
    total     = len(entries)
    flag_cats = {}
    for f in all_flags:
        flag_cats.setdefault(f.get("category", ""), []).append(f)
    no_auth   = sum(1 for e in entries if (e.auth_type or "") in ("None detected", "UNKNOWN", "none", ""))
    shadow    = sum(1 for e in entries if e.classification in ("Shadow", "Rogue"))
    param_eps = sum(1 for e in entries if re.search(r"/:\w+|/\{[\w]+\}", e.endpoint or ""))

    return [
        {"owasp_id": "API1",  "name": "Broken Object Level Authorization (BOLA)",
         "status": "REQUIRES TESTING", "affected_count": param_eps,
         "note": f"{param_eps} parameterized endpoints — manual pen-test required.",
         "conformance_level": "Level 0 — Not Tested"},
        {"owasp_id": "API2",  "name": "Broken Authentication",
         "status": "FAIL", "affected_count": no_auth,
         "note": f"{no_auth} of {total} endpoints have no detectable authentication.",
         "conformance_level": "Level 0 — Non-Conformant"},
        {"owasp_id": "API3",  "name": "Broken Object Property Level Authorization",
         "status": "NOT TESTED", "affected_count": 0,
         "note": "Requires runtime testing.", "conformance_level": "Level 0 — Not Tested"},
        {"owasp_id": "API4",  "name": "Unrestricted Resource Consumption",
         "status": "NOT TESTED", "affected_count": 0,
         "note": "Rate limiting not assessable without live traffic.",
         "conformance_level": "Level 0 — Not Tested"},
        {"owasp_id": "API5",  "name": "Broken Function Level Authorization",
         "status": "REQUIRES TESTING", "affected_count": len(flag_cats.get("API5", [])),
         "note": "Admin endpoints identified — role-based access validation required.",
         "conformance_level": "Level 0 — Not Tested"},
        {"owasp_id": "API6",  "name": "Unrestricted Access to Sensitive Business Flows",
         "status": "NOT TESTED", "affected_count": 0,
         "note": "Requires business logic review and runtime testing.",
         "conformance_level": "Level 0 — Not Tested"},
        {"owasp_id": "API7",  "name": "Server Side Request Forgery (SSRF)",
         "status": "NOT TESTED", "affected_count": 0,
         "note": "Active probing not performed in this engagement.",
         "conformance_level": "Level 0 — Not Tested"},
        {"owasp_id": "API8",  "name": "Security Misconfiguration",
         "status": "FAIL" if flag_cats.get("API8") else "PARTIAL",
         "affected_count": len(flag_cats.get("API8", [])),
         "note": f"{len(flag_cats.get('API8', []))} misconfiguration findings detected.",
         "conformance_level": "Level 1 — Partial"},
        {"owasp_id": "API9",  "name": "Improper Inventory Management",
         "status": "FAIL", "affected_count": shadow,
         "note": f"All {shadow} discovered endpoints are unregistered. No OpenAPI spec.",
         "conformance_level": "Level 0 — Non-Conformant"},
        {"owasp_id": "API10", "name": "Unsafe Consumption of APIs",
         "status": "REQUIRES REVIEW", "affected_count": len(flag_cats.get("API10", [])),
         "note": "External API integrations detected — consumption patterns require review.",
         "conformance_level": "Level 0 — Not Tested"},
    ]


# ---------------------------------------------------------------------------
# Inbound / Outbound summary builder
# ---------------------------------------------------------------------------

def _build_inbound_outbound_summary(inbound_entries: List[APIEntry],
                                     outbound_inventory: List[Dict]) -> Dict:
    """
    Builds the inbound vs outbound classification section.
    Inbound = the 503 APIs CRVM exposes (what others call INTO CRVM).
    Outbound = APIs CRVM calls OUT TO (external/internal services it depends on).
    """
    # Inbound breakdown
    inbound_by_sensitivity = _crit_counts(inbound_entries)
    inbound_by_method: Dict[str, int] = {}
    for e in inbound_entries:
        m = e.method or "UNKNOWN"
        inbound_by_method[m] = inbound_by_method.get(m, 0) + 1

    inbound_by_module: Dict[str, int] = {}
    for e in inbound_entries:
        mod = _infer_module(e)
        inbound_by_module[mod] = inbound_by_module.get(mod, 0) + 1

    # Outbound breakdown — strip internal _dedup_key before output
    clean_outbound = []
    for entry in outbound_inventory:
        clean = {k: v for k, v in entry.items() if not k.startswith("_")}
        clean_outbound.append(clean)

    outbound_by_exposure: Dict[str, int] = {}
    outbound_by_category: Dict[str, int] = {}
    outbound_by_auth: Dict[str, int]     = {}
    for entry in clean_outbound:
        exp = entry.get("exposure", "Unknown")
        cat = entry.get("category", "Unknown")
        auth = entry.get("auth_method", "unknown")
        outbound_by_exposure[exp]  = outbound_by_exposure.get(exp, 0) + 1
        outbound_by_category[cat]  = outbound_by_category.get(cat, 0) + 1
        outbound_by_auth[auth]     = outbound_by_auth.get(auth, 0) + 1

    hardcoded_key_count = sum(1 for e in clean_outbound if e.get("auth_method") == "hardcoded_key")
    external_outbound   = sum(1 for e in clean_outbound if e.get("exposure") == "External")
    internal_outbound   = sum(1 for e in clean_outbound if e.get("exposure") == "Internal")

    return {
        "inbound_apis": {
            "description":    "APIs that CRVM exposes — endpoints other systems or users call INTO CRVM",
            "total":          len(inbound_entries),
            "by_sensitivity": inbound_by_sensitivity,
            "by_method":      dict(sorted(inbound_by_method.items(), key=lambda x: -x[1])),
            "by_module":      dict(sorted(inbound_by_module.items(), key=lambda x: -x[1])[:10]),
            "all_classified_as": "Shadow — no API registry or OpenAPI spec exists",
        },
        "outbound_apis": {
            "description":           "APIs that CRVM calls OUT TO — external and internal services CRVM depends on",
            "total":                 len(clean_outbound),
            "external":              external_outbound,
            "internal":              internal_outbound,
            "hardcoded_credentials": hardcoded_key_count,
            "by_exposure":           outbound_by_exposure,
            "by_category":           outbound_by_category,
            "by_auth_method":        outbound_by_auth,
            "owasp_reference":       "API10 — Unsafe Consumption of Third-Party APIs",
            "risk_note": (
                f"{hardcoded_key_count} outbound API calls use hardcoded credentials — rotate immediately."
                if hardcoded_key_count > 0 else
                "No hardcoded credentials detected in outbound calls."
            ),
            "apis": clean_outbound,
        },
    }


# ---------------------------------------------------------------------------
# Executive summary
# ---------------------------------------------------------------------------

def _build_executive_summary(entries: List[APIEntry], counts: Dict,
                               secrets: List[Dict], outbound: List[Dict],
                               outbound_inventory: List[Dict] = None) -> Dict:
    total        = counts.get("total", len(entries))
    shadow_count = counts.get("Shadow", 0)
    rogue_count  = counts.get("Rogue", 0)
    crit_sens    = sum(1 for e in entries if e.data_sensitivity == "CRITICAL")
    high_risk    = sum(1 for e in entries if e.risk_score >= 50)
    secret_count = len(secrets)
    no_auth      = sum(1 for e in entries
                       if (e.auth_type or "") in ("None detected", "UNKNOWN", "none", ""))
    cve_count    = sum(len(e.cve_findings) for e in entries)
    outbound_total  = len(outbound_inventory or [])
    outbound_ext    = sum(1 for o in (outbound_inventory or []) if o.get("exposure") == "External")

    risk_level = "CRITICAL" if (crit_sens >= 10 or secret_count >= 5 or high_risk >= 100) else \
                 "HIGH"     if (high_risk >= 20  or secret_count >= 1) else "MEDIUM"

    secret_sentence = (
        f"Static analysis revealed {secret_count} hardcoded credential findings including "
        f"tokens, private keys, and environment variables. "
        if secret_count > 0 else
        "Secret scanning was performed; no hardcoded credentials were found. "
    )
    cve_sentence = (
        f"{cve_count} CVE findings were identified in application dependencies. "
        if cve_count > 0 else ""
    )
    outbound_sentence = (
        f"CRVM makes outbound calls to {outbound_total} external and internal APIs "
        f"({outbound_ext} external), none of which have been formally assessed for security. "
        if outbound_total > 0 else ""
    )

    narrative = (
        f"The API Discovery and Security Evaluation engagement for {counts.get('client','the client')} "
        f"identified {total} inbound APIs that CRVM exposes, and {outbound_total} outbound API dependencies "
        f"that CRVM calls. None of the {total} inbound endpoints were present in a formal API registry or "
        f"OpenAPI specification, resulting in all being classified as Shadow APIs — indicating a complete "
        f"absence of API governance controls. "
        f"{crit_sens} inbound endpoints handle CRITICAL sensitivity data and {high_risk} carry "
        f"High or Critical risk scores. "
        f"{secret_sentence}"
        f"{cve_sentence}"
        f"{outbound_sentence}"
        f"{no_auth} inbound endpoints show no detectable authentication mechanism. "
        f"The overall security posture for API governance is assessed as: {risk_level}."
    )

    top_recommendations = []
    priority = 1

    if secret_count > 0:
        top_recommendations.append({
            "priority": priority,
            "action": "Rotate all hardcoded credentials immediately",
            "rationale": f"{secret_count} secrets found in source code.",
            "effort": "Low", "impact": "CRITICAL",
        })
        priority += 1

    hardcoded_outbound = sum(1 for o in (outbound_inventory or []) if o.get("auth_method") == "hardcoded_key")
    if hardcoded_outbound > 0:
        top_recommendations.append({
            "priority": priority,
            "action": "Rotate hardcoded credentials in outbound API calls",
            "rationale": f"{hardcoded_outbound} outbound calls use hardcoded API keys.",
            "effort": "Low", "impact": "HIGH",
        })
        priority += 1

    if cve_count > 0:
        top_recommendations.append({
            "priority": priority,
            "action": "Patch vulnerable dependencies identified by CVE scan",
            "rationale": f"{cve_count} CVE findings in application packages.",
            "effort": "Medium", "impact": "HIGH",
        })
        priority += 1

    top_recommendations += [
        {
            "priority": priority,
            "action": "Create and enforce API inventory (OpenAPI spec + gateway registration)",
            "rationale": f"All {total} inbound endpoints are unregistered.",
            "effort": "High", "impact": "CRITICAL",
        },
        {
            "priority": priority + 1,
            "action": "Implement authentication middleware across all API routes",
            "rationale": f"{no_auth} inbound endpoints accessible without authentication.",
            "effort": "Medium", "impact": "HIGH",
        },
        {
            "priority": priority + 2,
            "action": "Deploy API gateway with rate limiting and security policies",
            "rationale": "No gateway currently in place — no centralized enforcement point.",
            "effort": "High", "impact": "HIGH",
        },
        {
            "priority": priority + 3,
            "action": "Assess and harden all outbound API integrations",
            "rationale": f"CRVM calls {outbound_total} external/internal APIs with no formal security assessment.",
            "effort": "Medium", "impact": "HIGH",
        },
    ]

    return {
        "overall_risk": risk_level,
        "narrative": narrative,
        "key_metrics": {
            "inbound_apis_total":              total,
            "outbound_apis_total":             outbound_total,
            "outbound_apis_external":          outbound_ext,
            "shadow_apis":                     shadow_count,
            "rogue_apis":                      rogue_count,
            "valid_apis":                      counts.get("Valid", 0),
            "critical_sensitivity_endpoints":  crit_sens,
            "high_critical_risk_endpoints":    high_risk,
            "hardcoded_secrets":               secret_count,
            "endpoints_without_auth":          no_auth,
            "external_integrations":           len(outbound),
            "cve_findings":                    cve_count,
        },
        "top_recommendations": top_recommendations,
    }


# ---------------------------------------------------------------------------
# Remediation
# ---------------------------------------------------------------------------

def _remediation_for(entry: APIEntry) -> str:
    cls   = entry.classification
    sens  = entry.data_sensitivity
    auth  = entry.auth_type or "None detected"
    score = entry.risk_score
    parts = []
    if cls in ("Shadow", "Rogue"):
        parts.append("Register in API gateway or decommission. Assign owner, add to OpenAPI spec.")
    if auth in ("None detected", "UNKNOWN", "none", ""):
        parts.append("Implement authentication middleware (JWT/OAuth2).")
    if sens == "CRITICAL":
        parts.append("Apply field-level encryption. Restrict to authorized roles only.")
    if score >= 75:
        parts.append("Immediate remediation required — escalate to security team.")
    if not parts:
        parts.append("Review as part of routine API governance cycle.")
    return " ".join(parts)


# ---------------------------------------------------------------------------
# Main Reporter class
# ---------------------------------------------------------------------------

class Reporter:
    def __init__(self, store, cfg: dict):
        self.store       = store
        self.cfg         = cfg
        out_cfg          = cfg.get("output", {})
        self.out_dir     = out_cfg.get("directory", "output")
        self.client_name = out_cfg.get("client_name", "Client")
        self.engagement  = out_cfg.get("engagement_name", "API Discovery & Security Evaluation")
        self.do_json     = out_cfg.get("json", True)
        self.do_pdf      = out_cfg.get("pdf", True)
        self.do_word     = out_cfg.get("word", True)
        self.scan_env    = out_cfg.get("scan_target_environment", "internal_non_prod")
        os.makedirs(self.out_dir, exist_ok=True)

    async def run(self):
        all_entries = self.store.all()
        counts      = self.store.count()
        counts["client"] = self.client_name

        clean_entries = [e for e in all_entries if not _is_noise(e)]
        for e in clean_entries:
            e.owasp_flags = _infer_owasp_flags(e)

        secrets_findings = getattr(self.store, "secrets_found", []) or []
        for s in secrets_findings:
            if "severity" not in s:
                s["severity"] = "CRITICAL"

        package_deps          = getattr(self.store, "package_dependencies", [])
        # NEW: rich outbound inventory from scanner
        outbound_api_inventory = getattr(self.store, "outbound_api_inventory", [])

        outbound_deps     = _extract_outbound_deps(all_entries)  # legacy compat
        tech_stack        = _detect_tech_stack(clean_entries, self.cfg)
        all_owasp_flags   = [f for e in clean_entries for f in e.owasp_flags]
        owasp_conformance = _build_owasp_conformance(clean_entries, all_owasp_flags)

        # Build inbound/outbound summary
        inbound_outbound = _build_inbound_outbound_summary(clean_entries, outbound_api_inventory)

        exec_summary = _build_executive_summary(
            clean_entries, counts, secrets_findings, outbound_deps, outbound_api_inventory
        )

        print(f"    Secrets loaded from store: {len(secrets_findings)} findings")
        print(f"    Package deps loaded: {len(package_deps)} packages")
        print(f"    Outbound APIs loaded: {len(outbound_api_inventory)} entries")

        if self.do_json:
            self._write_json(clean_entries, counts, secrets_findings, package_deps,
                             outbound_deps, tech_stack, owasp_conformance,
                             exec_summary, all_owasp_flags, inbound_outbound)
        if self.do_pdf:
            self._write_pdf(clean_entries, counts, secrets_findings,
                            outbound_deps, owasp_conformance, exec_summary,
                            all_owasp_flags, inbound_outbound)
        if self.do_word:
            self._write_word(clean_entries, counts, secrets_findings,
                             outbound_deps, owasp_conformance, exec_summary,
                             all_owasp_flags, inbound_outbound)

        self._print_summary(counts, clean_entries, secrets_findings, outbound_api_inventory)

    # ------------------------------------------------------------------
    # JSON output
    # ------------------------------------------------------------------

    def _write_json(self, entries, counts, secrets, package_deps, outbound, tech_stack,
                    owasp_conformance, exec_summary, all_owasp_flags, inbound_outbound):

        def _enrich(e: APIEntry) -> Dict:
            d = e.to_dict()
            d["functional_module"] = _infer_module(e)
            d["api_version"]       = _infer_version(e.endpoint)
            d["source_file"]       = (e.evidence or {}).get("file", "unknown")
            d["inferred_owner"]    = e.inferred_owner or "Pending Triage - Application Owner"
            d["risk_band"]         = _risk_band(e.risk_score)
            d["owasp_flags"]       = e.owasp_flags
            d["cve_findings"]      = e.cve_findings
            d["remediation"]       = _remediation_for(e)
            return d

        by_class = {
            cls: [_enrich(e) for e in entries if e.classification == cls]
            for cls in ["Valid", "Shadow", "New", "Rogue", "UNCLASSIFIED"]
        }

        shadow_rogue = sorted(
            [_enrich(e) for e in entries if e.classification in ("Shadow", "Rogue")],
            key=lambda x: x.get("risk_score", 0), reverse=True
        )

        owasp_findings_list = sorted(
            all_owasp_flags,
            key=lambda x: {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}.get(x.get("severity", "LOW"), 3)
        )

        bom = self._build_bom(entries, tech_stack, outbound, package_deps)

        all_cves  = []
        seen_cves = set()
        for e in entries:
            for cve in e.cve_findings:
                cid = cve.get("cve", "")
                if cid not in seen_cves:
                    all_cves.append({**cve, "endpoint_count": sum(
                        1 for x in entries if any(c.get("cve") == cid for c in x.cve_findings)
                    )})
                    seen_cves.add(cid)
        all_cves.sort(key=lambda x: x.get("cvss", 0), reverse=True)

        output = {
            "schema_version":           "2.2",
            "engagement":               self.engagement,
            "client":                   self.client_name,
            "scan_target_environment":  self.scan_env,
            "generated_at":             datetime.utcnow().isoformat() + "Z",

            "executive_summary": exec_summary,

            "summary": {
                **{k: v for k, v in counts.items() if k != "client"},
                "api_inventory_count":       len(entries),
                "noise_filtered_count":      counts.get("total", 0) - len(entries),
                "inbound_api_count":         len(entries),
                "outbound_api_count":        len(inbound_outbound["outbound_apis"]["apis"]),
                "outbound_external_count":   inbound_outbound["outbound_apis"]["external"],
                "outbound_internal_count":   inbound_outbound["outbound_apis"]["internal"],
                "secrets_count":             len(secrets),
                "external_integrations":     len(outbound),
                "data_sensitivity":          _crit_counts(entries),
                "owasp_findings_total":      len(owasp_findings_list),
                "inferred_owasp_findings":   sum(1 for f in owasp_findings_list if f.get("source") == "inferred"),
                "live_owasp_findings":       sum(1 for f in owasp_findings_list if f.get("source") != "inferred"),
                "cve_findings_total":        len(all_cves),
                "high_critical_risk_count":  sum(1 for e in entries if e.risk_score >= 50),
                "endpoints_without_auth":    sum(
                    1 for e in entries
                    if (e.auth_type or "") in ("None detected", "UNKNOWN", "none", "")
                ),
            },

            # Dedicated inbound/outbound classification section — NEW
            "inbound_outbound_classification": inbound_outbound,

            "api_inventory":             {cls: apis for cls, apis in by_class.items() if apis},
            "shadow_rogue_register":     shadow_rogue,
            "owasp_findings":            owasp_findings_list,
            "owasp_conformance_summary": owasp_conformance,
            "secrets_findings":          secrets,
            "cve_findings_summary":      all_cves,
            "outbound_dependencies":     outbound,
            "api_bom":                   bom,

            "all_endpoints": [
                _enrich(e)
                for e in sorted(entries, key=lambda x: x.risk_score, reverse=True)
            ],
        }

        path = os.path.join(self.out_dir, "api_discovery_full.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2)
        print(f"    JSON: {path}")

        reg_path = os.path.join(self.out_dir, "shadow_rogue_register.json")
        with open(reg_path, "w", encoding="utf-8") as f:
            json.dump({
                "generated_at": datetime.utcnow().isoformat() + "Z",
                "engagement":   self.engagement,
                "client":       self.client_name,
                "total":        len(shadow_rogue),
                "shadow_count": sum(1 for e in shadow_rogue if e["classification"] == "Shadow"),
                "rogue_count":  sum(1 for e in shadow_rogue if e["classification"] == "Rogue"),
                "shadow": [e for e in shadow_rogue if e["classification"] == "Shadow"],
                "rogue":  [e for e in shadow_rogue if e["classification"] == "Rogue"],
            }, f, indent=2)
        print(f"    JSON: {reg_path}")

        sec_path = os.path.join(self.out_dir, "secrets_findings.json")
        with open(sec_path, "w", encoding="utf-8") as f:
            json.dump({"generated_at": datetime.utcnow().isoformat() + "Z",
                       "total": len(secrets), "findings": secrets}, f, indent=2)
        print(f"    JSON: {sec_path}")

        # NEW: dedicated outbound inventory file
        out_path = os.path.join(self.out_dir, "outbound_api_inventory.json")
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump({
                "generated_at": datetime.utcnow().isoformat() + "Z",
                "engagement":   self.engagement,
                "client":       self.client_name,
                "total":        len(inbound_outbound["outbound_apis"]["apis"]),
                "external":     inbound_outbound["outbound_apis"]["external"],
                "internal":     inbound_outbound["outbound_apis"]["internal"],
                "apis":         inbound_outbound["outbound_apis"]["apis"],
            }, f, indent=2)
        print(f"    JSON: {out_path}")

    def _build_bom(self, entries, tech_stack, outbound, package_deps=None) -> Dict:
        items = []
        for e in entries:
            items.append({
                "endpoint":          e.endpoint,
                "method":            e.method,
                "classification":    e.classification,
                "functional_module": _infer_module(e),
                "api_version":       _infer_version(e.endpoint),
                "auth_type":         e.auth_type,
                "data_sensitivity":  e.data_sensitivity,
                "exposure":          e.exposure,
                "environment":       e.environment or self.scan_env,
                "owner":             e.inferred_owner or e.owner or "Pending review",
                "risk_score":        e.risk_score,
                "risk_band":         _risk_band(e.risk_score),
                "owasp_categories":  [f.get("category") for f in e.owasp_flags],
                "cve_count":         len(e.cve_findings),
                "discovered_by":     e.discovered_by,
                "source_file":       (e.evidence or {}).get("file", "unknown"),
                "first_seen":        e.first_seen,
                "last_seen":         e.last_seen,
                "tags":              e.tags,
            })
        items.sort(key=lambda x: x["risk_score"], reverse=True)
        return {
            "tech_stack":   tech_stack,
            "upstream_downstream_dependencies": outbound,
            "package_dependencies": package_deps or [],
            "api_endpoints": items,
        }

    # ------------------------------------------------------------------
    # PDF output
    # ------------------------------------------------------------------

    def _write_pdf(self, entries, counts, secrets, outbound,
                   owasp_conformance, exec_summary, all_owasp_flags, inbound_outbound):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                             Table, TableStyle, PageBreak, HRFlowable)

            path = os.path.join(self.out_dir, "executive_report.pdf")
            doc  = SimpleDocTemplate(path, pagesize=A4,
                                     leftMargin=2*cm, rightMargin=2*cm,
                                     topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()

            title_s = ParagraphStyle("TitleS", parent=styles["Title"],
                                      fontSize=22, textColor=colors.HexColor("#1a2744"), spaceAfter=4)
            sub_s   = ParagraphStyle("SubS",   parent=styles["Normal"],
                                      fontSize=11, textColor=colors.HexColor("#555555"), spaceAfter=4)
            h1_s    = ParagraphStyle("H1S",    parent=styles["Heading1"],
                                      fontSize=16, textColor=colors.HexColor("#1a2744"),
                                      spaceBefore=14, spaceAfter=6)
            h2_s    = ParagraphStyle("H2S",    parent=styles["Heading2"],
                                      fontSize=13, textColor=colors.HexColor("#1a2744"),
                                      spaceBefore=10, spaceAfter=4)
            body_s  = ParagraphStyle("BodyS",  parent=styles["Normal"], fontSize=9, leading=13)

            DARK    = colors.HexColor("#1a2744")
            RED     = colors.HexColor("#c0392b")
            ORANGE  = colors.HexColor("#e67e22")
            GREEN   = colors.HexColor("#27ae60")
            TEAL    = colors.HexColor("#16a085")
            LGREY   = colors.HexColor("#f5f5f5")
            LRED    = colors.HexColor("#fff5f5")
            LORANGE = colors.HexColor("#fef9f0")
            LTEAL   = colors.HexColor("#f0fafa")
            WHITE   = colors.white
            sev_colors = {"CRITICAL": RED, "HIGH": ORANGE,
                          "MEDIUM": colors.HexColor("#f39c12"), "LOW": GREEN}

            def tbl(data, col_widths, hdr_color=DARK, alt=LGREY):
                t = Table(data, colWidths=col_widths, repeatRows=1)
                t.setStyle(TableStyle([
                    ("BACKGROUND",     (0, 0), (-1, 0),  hdr_color),
                    ("TEXTCOLOR",      (0, 0), (-1, 0),  WHITE),
                    ("FONTNAME",       (0, 0), (-1, 0),  "Helvetica-Bold"),
                    ("FONTSIZE",       (0, 0), (-1, -1), 8),
                    ("PADDING",        (0, 0), (-1, -1), 5),
                    ("GRID",           (0, 0), (-1, -1), 0.4, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, alt]),
                    ("VALIGN",         (0, 0), (-1, -1), "TOP"),
                ]))
                return t

            story = []

            # ---- Cover ----
            story.append(Spacer(1, 1*cm))
            story.append(Paragraph(self.engagement, title_s))
            story.append(Paragraph(f"Client: <b>{self.client_name}</b>", sub_s))
            story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%d %B %Y %H:%M UTC')}", sub_s))
            story.append(Paragraph(f"Environment: {self.scan_env}", sub_s))
            story.append(HRFlowable(width="100%", thickness=2, color=DARK))
            story.append(Spacer(1, 0.4*cm))

            risk       = exec_summary.get("overall_risk", "UNKNOWN")
            risk_color = sev_colors.get(risk, colors.grey)
            rt = Table([["OVERALL RISK ASSESSMENT", risk]], colWidths=[12*cm, 4*cm])
            rt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, 0), DARK),
                ("BACKGROUND", (1, 0), (1, 0), risk_color),
                ("TEXTCOLOR",  (0, 0), (-1, -1), WHITE),
                ("FONTNAME",   (0, 0), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 12),
                ("PADDING",    (0, 0), (-1, -1), 8),
                ("ALIGN",      (1, 0), (1, 0), "CENTER"),
            ]))
            story.append(rt)
            story.append(Spacer(1, 0.5*cm))

            # ---- 1. Executive Summary ----
            story.append(Paragraph("1. Executive Summary", h1_s))
            story.append(Paragraph(exec_summary.get("narrative", ""), body_s))
            story.append(Spacer(1, 0.4*cm))

            km = exec_summary.get("key_metrics", {})
            metrics_data = [
                ["Metric", "Value", "Metric", "Value"],
                ["Inbound APIs (exposed)",    str(km.get("inbound_apis_total", 0)),
                 "Outbound APIs (called)",    str(km.get("outbound_apis_total", 0))],
                ["Shadow APIs",               str(km.get("shadow_apis", 0)),
                 "Valid APIs",                str(km.get("valid_apis", 0))],
                ["Critical Sensitivity",      str(km.get("critical_sensitivity_endpoints", 0)),
                 "High/Critical Risk",        str(km.get("high_critical_risk_endpoints", 0))],
                ["Hardcoded Secrets",         str(km.get("hardcoded_secrets", 0)),
                 "No Auth Detected",          str(km.get("endpoints_without_auth", 0))],
                ["CVE Findings",              str(km.get("cve_findings", 0)),
                 "Outbound External APIs",    str(km.get("outbound_apis_external", 0))],
            ]
            story.append(tbl(metrics_data, [5*cm, 3*cm, 5*cm, 3*cm]))
            story.append(Spacer(1, 0.4*cm))

            story.append(Paragraph("Top Recommendations", h2_s))
            rec_data = [["#", "Action", "Impact", "Effort"]]
            for r in exec_summary.get("top_recommendations", []):
                rec_data.append([str(r.get("priority","")), r.get("action",""),
                                  r.get("impact",""), r.get("effort","")])
            story.append(tbl(rec_data, [0.8*cm, 10*cm, 2.5*cm, 2.5*cm]))
            story.append(PageBreak())

            # ---- 2. API Classification Summary ----
            story.append(Paragraph("2. API Classification Summary", h1_s))
            cls_data = [
                ["Classification", "Count", "Description"],
                ["Valid",  str(counts.get("Valid",  0)), "Documented and authorized APIs"],
                ["Shadow", str(counts.get("Shadow", 0)), "Undocumented — present in source, not in registry"],
                ["New",    str(counts.get("New",    0)), "Recently discovered, not yet catalogued"],
                ["Rogue",  str(counts.get("Rogue",  0)), "Unauthorized, outside governance scope"],
            ]
            story.append(tbl(cls_data, [4*cm, 3*cm, 10*cm]))
            story.append(Spacer(1, 0.4*cm))

            story.append(Paragraph("Data Sensitivity Distribution", h2_s))
            sens      = _crit_counts(entries)
            sens_data = [["Sensitivity", "Count"]] + [[lvl, str(cnt)] for lvl, cnt in sens.items() if cnt]
            story.append(tbl(sens_data, [6*cm, 6*cm]))
            story.append(Spacer(1, 0.4*cm))

            story.append(Paragraph("Endpoints by Functional Module", h2_s))
            module_counts: Dict[str, int] = {}
            for e in entries:
                m = _infer_module(e)
                module_counts[m] = module_counts.get(m, 0) + 1
            mod_data = [["Functional Module", "Count"]] + sorted(
                [[k, str(v)] for k, v in module_counts.items()], key=lambda x: -int(x[1]))
            story.append(tbl(mod_data, [12*cm, 4*cm]))
            story.append(PageBreak())

            # ---- 3. Top Risk Endpoints ----
            story.append(Paragraph("3. Top 25 Highest Risk Endpoints", h1_s))
            top       = sorted(entries, key=lambda e: e.risk_score, reverse=True)[:25]
            risk_data = [["Endpoint", "Module", "Method", "Risk", "Sensitivity"]]
            for e in top:
                ep = (e.endpoint[:55] + "…") if len(e.endpoint) > 55 else e.endpoint
                risk_data.append([ep, _infer_module(e), e.method or "?",
                                   f"{e.risk_score} ({_risk_band(e.risk_score)})", e.data_sensitivity])
            story.append(tbl(risk_data, [7*cm, 3.5*cm, 1.5*cm, 3*cm, 2*cm], hdr_color=RED, alt=LRED))
            story.append(PageBreak())

            # ---- 4. OWASP Conformance ----
            story.append(Paragraph("4. OWASP API Top 10 Conformance Summary", h1_s))
            conf_data = [["ID", "Category", "Status", "Affected", "Conformance"]]
            for row in owasp_conformance:
                conf_data.append([row["owasp_id"], row["name"], row["status"],
                                   str(row["affected_count"]), row["conformance_level"]])
            story.append(tbl(conf_data, [1.5*cm, 5.5*cm, 2.5*cm, 2*cm, 5*cm],
                             hdr_color=ORANGE, alt=LORANGE))
            story.append(Spacer(1, 0.4*cm))

            if all_owasp_flags:
                story.append(Paragraph("OWASP Findings Detail (Top 40)", h2_s))
                owasp_sorted = sorted(
                    all_owasp_flags,
                    key=lambda x: {"CRITICAL":0,"HIGH":1,"MEDIUM":2,"LOW":3}.get(x.get("severity","LOW"),3)
                )[:40]
                owasp_det = [["Category", "Finding", "Sev", "Source", "Endpoint"]]
                for f in owasp_sorted:
                    ep = (f.get("endpoint","")[:35]+"…") if len(f.get("endpoint",""))>35 else f.get("endpoint","")
                    owasp_det.append([
                        f.get("category",""),
                        (f.get("finding","")[:55]+"…") if len(f.get("finding",""))>55 else f.get("finding",""),
                        f.get("severity",""), f.get("source","live"), ep,
                    ])
                story.append(tbl(owasp_det, [1.8*cm, 6.5*cm, 1.8*cm, 1.5*cm, 5*cm],
                                 hdr_color=ORANGE, alt=LORANGE))
            story.append(PageBreak())

            # ---- 5. Secrets ----
            if secrets:
                story.append(Paragraph("5. Hardcoded Secrets & Credentials", h1_s))
                story.append(Paragraph(
                    f"<font color='red'><b>WARNING:</b></font> "
                    f"{len(secrets)} hardcoded credential(s) detected. Rotate all findings immediately.",
                    body_s))
                story.append(Spacer(1, 0.3*cm))
                sec_data = [["Type", "Severity", "File", "Recommendation"]]
                for s in secrets:
                    f_short = (s["file"][-50:]+"…") if len(s.get("file",""))>50 else s.get("file","")
                    sec_data.append([s.get("type",""), s.get("severity","CRITICAL"),
                                     f_short, s.get("recommendation","")[:60]])
                story.append(tbl(sec_data, [2.5*cm, 2*cm, 6*cm, 7*cm], hdr_color=RED, alt=LRED))
                story.append(PageBreak())

            # ---- 6. Outbound Dependencies (legacy) ----
            if outbound:
                story.append(Paragraph("6. Known Outbound API Dependencies", h1_s))
                dep_data = [["Integration", "Category", "Exposure", "Risk", "Recommendation"]]
                for d in outbound:
                    dep_data.append([d["integration"], d["category"], d["exposure"],
                                     d["risk"], d["recommendation"][:60]])
                story.append(tbl(dep_data, [3*cm, 3.5*cm, 2*cm, 2*cm, 7*cm]))
                story.append(PageBreak())

            # ---- 7. Shadow & Rogue Register ----
            shadow_rogue = [e for e in entries if e.classification in ("Shadow", "Rogue")]
            if shadow_rogue:
                story.append(Paragraph("7. Shadow & Rogue API Register", h1_s))
                sr_data = [["Endpoint", "Module", "Type", "Risk", "Sensitivity", "Method", "Owner"]]
                for e in sorted(shadow_rogue, key=lambda x: x.risk_score, reverse=True)[:60]:
                    ep    = (e.endpoint[:45]+"…") if len(e.endpoint)>45 else e.endpoint
                    owner = (e.inferred_owner or "Pending Triage")[:30]
                    sr_data.append([ep, _infer_module(e), e.classification,
                                    f"{e.risk_score} ({_risk_band(e.risk_score)})",
                                    e.data_sensitivity, e.method or "?", owner])
                story.append(tbl(sr_data,
                                 [5.5*cm, 3*cm, 1.8*cm, 2.2*cm, 1.8*cm, 1.2*cm, 3*cm],
                                 hdr_color=DARK, alt=LGREY))
                story.append(PageBreak())

            # ---- 8. Inbound vs Outbound API Classification ---- NEW
            story.append(Paragraph("8. Inbound vs Outbound API Classification", h1_s))

            # Summary box
            ib    = inbound_outbound["inbound_apis"]
            ob    = inbound_outbound["outbound_apis"]
            io_summary = [
                ["Direction", "Total", "External", "Internal", "OWASP Reference"],
                ["Inbound (APIs CRVM exposes)",  str(ib["total"]), "N/A", "All internal network", "API2, API9"],
                ["Outbound (APIs CRVM calls)",   str(ob["total"]), str(ob["external"]),
                 str(ob["internal"]), "API10"],
            ]
            story.append(tbl(io_summary, [5.5*cm, 2*cm, 2.5*cm, 3.5*cm, 3.5*cm],
                             hdr_color=TEAL, alt=LTEAL))
            story.append(Spacer(1, 0.4*cm))

            # Outbound inventory table
            outbound_apis = ob.get("apis", [])
            if outbound_apis:
                story.append(Paragraph("Outbound API Inventory", h2_s))
                if ob.get("hardcoded_credentials", 0) > 0:
                    story.append(Paragraph(
                        f"<font color='red'><b>WARNING:</b></font> "
                        f"{ob['hardcoded_credentials']} outbound call(s) use hardcoded credentials. "
                        "Rotate immediately.",
                        body_s))
                    story.append(Spacer(1, 0.2*cm))

                out_data = [["URL / Integration", "Category", "Exposure", "Method", "Auth Method", "Risk"]]
                for o in sorted(outbound_apis, key=lambda x: {"HIGH":0,"MEDIUM":1,"LOW":2}.get(x.get("risk","LOW"),2)):
                    url_short = (o.get("url","")[:45]+"…") if len(o.get("url",""))>45 else o.get("url","")
                    out_data.append([
                        url_short,
                        o.get("category",""),
                        o.get("exposure",""),
                        o.get("method","UNKNOWN"),
                        o.get("auth_method","unknown"),
                        o.get("risk",""),
                    ])
                story.append(tbl(out_data, [5.5*cm, 3*cm, 2*cm, 1.8*cm, 2.5*cm, 2*cm],
                                 hdr_color=TEAL, alt=LTEAL))
                story.append(Spacer(1, 0.4*cm))

                # Auth method breakdown
                story.append(Paragraph("Outbound Auth Method Breakdown", h2_s))
                auth_data = [["Auth Method", "Count", "Risk Implication"]]
                auth_implications = {
                    "hardcoded_key":  "CRITICAL — rotate immediately",
                    "env_variable":   "LOW — acceptable if secrets manager backed",
                    "bearer_token":   "MEDIUM — verify token rotation policy",
                    "basic_auth":     "MEDIUM — consider upgrading to token-based",
                    "none":           "HIGH — verify if endpoint is truly public",
                    "unknown":        "MEDIUM — manual review required",
                }
                for auth, count in sorted(ob.get("by_auth_method",{}).items(), key=lambda x: -x[1]):
                    auth_data.append([auth, str(count), auth_implications.get(auth, "Review required")])
                story.append(tbl(auth_data, [5*cm, 2*cm, 10*cm], hdr_color=TEAL, alt=LTEAL))

            doc.build(story)
            print(f"    PDF: {path}")

        except ImportError:
            print("    reportlab required: pip install reportlab")
        except Exception as exc:
            import traceback
            print(f"    PDF error: {exc}")
            traceback.print_exc()

    # ------------------------------------------------------------------
    # Word output
    # ------------------------------------------------------------------

    def _write_word(self, entries, counts, secrets, outbound,
                    owasp_conformance, exec_summary, all_owasp_flags, inbound_outbound):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            def add_table(doc, headers, rows, bold_header=True):
                t = doc.add_table(rows=1, cols=len(headers))
                t.style = "Table Grid"
                hdr = t.rows[0].cells
                for i, h in enumerate(headers):
                    hdr[i].text = h
                    if bold_header and hdr[i].paragraphs[0].runs:
                        hdr[i].paragraphs[0].runs[0].bold = True
                    elif bold_header:
                        hdr[i].paragraphs[0].add_run().bold = True
                for row_data in rows:
                    r = t.add_row().cells
                    for i, val in enumerate(row_data):
                        r[i].text = str(val)
                return t

            # ---- Cover ----
            title = doc.add_heading(self.engagement, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Client: {self.client_name}")
            doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%d %B %Y %H:%M UTC')}")
            doc.add_paragraph(f"Environment: {self.scan_env}")
            doc.add_paragraph(f"Overall Risk: {exec_summary.get('overall_risk', 'UNKNOWN')}")
            doc.add_page_break()

            # ---- 1. Executive Summary ----
            doc.add_heading("1. Executive Summary", 1)
            doc.add_paragraph(exec_summary.get("narrative", ""))
            doc.add_heading("Key Metrics", 2)
            km = exec_summary.get("key_metrics", {})
            metrics = [
                ("Inbound APIs (CRVM exposes)",   km.get("inbound_apis_total", 0)),
                ("Outbound APIs (CRVM calls)",    km.get("outbound_apis_total", 0)),
                ("Outbound External APIs",        km.get("outbound_apis_external", 0)),
                ("Shadow APIs",                   km.get("shadow_apis", 0)),
                ("Valid APIs",                    km.get("valid_apis", 0)),
                ("Critical Sensitivity Endpoints",km.get("critical_sensitivity_endpoints", 0)),
                ("High/Critical Risk Endpoints",  km.get("high_critical_risk_endpoints", 0)),
                ("Hardcoded Secrets Found",       km.get("hardcoded_secrets", 0)),
                ("Endpoints Without Auth",        km.get("endpoints_without_auth", 0)),
                ("CVE Findings",                  km.get("cve_findings", 0)),
            ]
            add_table(doc, ["Metric", "Value"], [[k, str(v)] for k, v in metrics])
            doc.add_heading("Top Recommendations", 2)
            add_table(doc, ["#", "Action", "Impact", "Effort"],
                      [[str(r["priority"]), r["action"], r["impact"], r["effort"]]
                       for r in exec_summary.get("top_recommendations", [])])
            doc.add_page_break()

            # ---- 2. Classification ----
            doc.add_heading("2. API Classification Summary", 1)
            add_table(doc, ["Classification", "Count", "Description"], [
                ("Valid",  counts.get("Valid",  0), "Documented and authorized"),
                ("Shadow", counts.get("Shadow", 0), "Undocumented, present in source"),
                ("New",    counts.get("New",    0), "Recently appeared"),
                ("Rogue",  counts.get("Rogue",  0), "Unauthorized, outside governance"),
            ])
            doc.add_paragraph("")
            doc.add_heading("Data Sensitivity", 2)
            sens = _crit_counts(entries)
            add_table(doc, ["Sensitivity Level", "Count"], [[k, str(v)] for k, v in sens.items() if v])
            doc.add_heading("Endpoints by Functional Module", 2)
            module_counts: Dict[str, int] = {}
            for e in entries:
                m = _infer_module(e)
                module_counts[m] = module_counts.get(m, 0) + 1
            add_table(doc, ["Functional Module", "Count"],
                      sorted([[k, str(v)] for k, v in module_counts.items()], key=lambda x: -int(x[1])))
            doc.add_page_break()

            # ---- 3. Top Risk ----
            doc.add_heading("3. Top 25 Highest Risk Endpoints", 1)
            top = sorted(entries, key=lambda e: e.risk_score, reverse=True)[:25]
            add_table(doc, ["Endpoint", "Module", "Method", "Risk Score", "Sensitivity"],
                      [[e.endpoint[:80], _infer_module(e), e.method or "?",
                        f"{e.risk_score} ({_risk_band(e.risk_score)})", e.data_sensitivity]
                       for e in top])
            doc.add_page_break()

            # ---- 4. OWASP ----
            doc.add_heading("4. OWASP API Top 10 Conformance Summary", 1)
            add_table(doc, ["ID", "Category", "Status", "Affected Count", "Conformance Level"],
                      [[r["owasp_id"], r["name"], r["status"],
                        str(r["affected_count"]), r["conformance_level"]]
                       for r in owasp_conformance])
            if all_owasp_flags:
                doc.add_heading("OWASP Findings (Top 50)", 2)
                owasp_sorted = sorted(
                    all_owasp_flags,
                    key=lambda x: {"CRITICAL":0,"HIGH":1,"MEDIUM":2,"LOW":3}.get(x.get("severity","LOW"),3)
                )[:50]
                add_table(doc, ["Category", "Finding", "Severity", "Source", "Remediation"],
                          [[f.get("category",""), f.get("finding","")[:80],
                            f.get("severity",""), f.get("source","live"),
                            f.get("remediation","")[:100]]
                           for f in owasp_sorted])
            doc.add_page_break()

            # ---- 5. Secrets ----
            if secrets:
                doc.add_heading("5. Hardcoded Secrets & Credentials", 1)
                doc.add_paragraph(
                    f"WARNING: {len(secrets)} hardcoded credentials detected. "
                    "Rotate immediately and migrate to a secrets manager."
                )
                add_table(doc, ["Type", "Severity", "File", "Preview", "Recommendation"],
                          [[s.get("type",""), s.get("severity","CRITICAL"),
                            s.get("file","")[-60:], s.get("match_preview","")[:40],
                            s.get("recommendation","")[:80]]
                           for s in secrets])
                doc.add_page_break()

            # ---- 6. Outbound Dependencies ----
            if outbound:
                doc.add_heading("6. Known Outbound API Dependencies", 1)
                add_table(doc, ["Integration", "Category", "Exposure", "Risk", "Recommendation"],
                          [[d["integration"], d["category"], d["exposure"],
                            d["risk"], d["recommendation"][:80]]
                           for d in outbound])
                doc.add_page_break()

            # ---- 7. Shadow & Rogue Register ----
            shadow_rogue = [e for e in entries if e.classification in ("Shadow", "Rogue")]
            if shadow_rogue:
                doc.add_heading("7. Shadow & Rogue API Register", 1)
                add_table(doc,
                          ["Endpoint", "Module", "Type", "Risk Score", "Sensitivity", "Method", "Owner", "Discovered By"],
                          [[e.endpoint[:80], _infer_module(e), e.classification,
                            f"{e.risk_score} ({_risk_band(e.risk_score)})",
                            e.data_sensitivity, e.method or "?",
                            e.inferred_owner or "Pending Triage",
                            ", ".join(e.discovered_by[:3])]
                           for e in sorted(shadow_rogue, key=lambda x: x.risk_score, reverse=True)])
                doc.add_page_break()

            # ---- 8. Inbound vs Outbound — NEW ----
            doc.add_heading("8. Inbound vs Outbound API Classification", 1)
            ib = inbound_outbound["inbound_apis"]
            ob = inbound_outbound["outbound_apis"]

            doc.add_paragraph(
                f"CRVM exposes {ib['total']} inbound APIs (endpoints other systems call INTO CRVM) "
                f"and makes outbound calls to {ob['total']} external and internal APIs "
                f"({ob['external']} external, {ob['internal']} internal)."
            )

            doc.add_heading("Inbound vs Outbound Summary", 2)
            add_table(doc, ["Direction", "Total", "External", "Internal", "OWASP Reference"], [
                ("Inbound — APIs CRVM exposes",  ib["total"], "N/A", "All internal network", "API2, API9"),
                ("Outbound — APIs CRVM calls",   ob["total"], ob["external"], ob["internal"], "API10"),
            ])
            doc.add_paragraph("")

            # Outbound inventory
            outbound_apis = ob.get("apis", [])
            if outbound_apis:
                doc.add_heading("Outbound API Inventory", 2)
                if ob.get("hardcoded_credentials", 0) > 0:
                    doc.add_paragraph(
                        f"WARNING: {ob['hardcoded_credentials']} outbound call(s) use hardcoded "
                        "credentials — rotate immediately."
                    )
                add_table(doc,
                          ["URL / Host", "Integration", "Category", "Exposure", "Method", "Auth Method", "Risk", "Source Files"],
                          [[o.get("url","")[:60], o.get("integration",""), o.get("category",""),
                            o.get("exposure",""), o.get("method","UNKNOWN"),
                            o.get("auth_method","unknown"), o.get("risk",""),
                            ", ".join(o.get("source_files",[])[:2])]
                           for o in sorted(outbound_apis,
                                           key=lambda x: {"HIGH":0,"MEDIUM":1,"LOW":2}.get(x.get("risk","LOW"),2))])
                doc.add_paragraph("")

                doc.add_heading("Outbound Auth Method Breakdown", 2)
                auth_implications = {
                    "hardcoded_key":  "CRITICAL — rotate immediately",
                    "env_variable":   "LOW — acceptable if secrets manager backed",
                    "bearer_token":   "MEDIUM — verify token rotation policy",
                    "basic_auth":     "MEDIUM — consider upgrading to token-based",
                    "none":           "HIGH — verify if endpoint is truly public",
                    "unknown":        "MEDIUM — manual review required",
                }
                add_table(doc, ["Auth Method", "Count", "Risk Implication"],
                          [[auth, str(count), auth_implications.get(auth, "Review required")]
                           for auth, count in sorted(ob.get("by_auth_method",{}).items(), key=lambda x: -x[1])])

            path = os.path.join(self.out_dir, "api_discovery_report.docx")
            doc.save(path)
            print(f"    Word: {path}")

        except ImportError:
            print("    python-docx required: pip install python-docx")
        except Exception as exc:
            import traceback
            print(f"    Word error: {exc}")
            traceback.print_exc()

    # ------------------------------------------------------------------
    # Console summary
    # ------------------------------------------------------------------

    def _print_summary(self, counts, entries, secrets, outbound_inventory=None):
        high_risk      = sum(1 for e in entries if e.risk_score >= 50)
        owasp_count    = sum(len(e.owasp_flags) for e in entries)
        cve_count      = sum(len(e.cve_findings) for e in entries)
        outbound_count = len(outbound_inventory or [])
        outbound_ext   = sum(1 for o in (outbound_inventory or []) if o.get("exposure") == "External")

        print(f"\n{'='*70}")
        print(f"  {self.engagement}")
        print(f"  Client  : {self.client_name}")
        print(f"  Env     : {self.scan_env}")
        print(f"{'='*70}")
        print(f"  Inbound APIs (exposed by CRVM) : {len(entries)}")
        print(f"  Outbound APIs (called by CRVM) : {outbound_count} ({outbound_ext} external)")
        print(f"  Valid                          : {counts.get('Valid', 0)}")
        print(f"  Shadow                         : {counts.get('Shadow', 0)}")
        print(f"  New                            : {counts.get('New', 0)}")
        print(f"  Rogue                          : {counts.get('Rogue', 0)}")
        print(f"  High/Critical risk             : {high_risk}")
        print(f"  OWASP findings (incl.inferred) : {owasp_count}")
        print(f"  Hardcoded secrets              : {len(secrets)}")
        print(f"  CVE findings                   : {cve_count}")
        print(f"  Output directory               : {os.path.abspath(self.out_dir)}")
        print(f"{'='*70}\n")